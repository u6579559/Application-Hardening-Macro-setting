import os

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches
from docx.shared import RGBColor
from docx.shared import Pt


def create_essential_eight_document(data, filepaths, type):
    """Types:
        1: Application Only
        2: Macro Only
        3: Both
    """
    document = Document()
    styles = document.styles
    create_fonts(document, styles)
    filepathString = "";

    for filepath in filepaths:
        filepath = filepath[filepath.rfind('\\') + 1:len(filepath)]
        file_name = os.path.basename(filepath)
        filepathString = filepathString + "\n" + file_name

    if type == 1:
        create_document_header(document, type, data)
        create_application_framework(document)
        populate_application_framework(document, data)
    elif type == 2:
        create_document_header(document, type, data)
        create_macro_framework(document)
        populate_macro_framework(document, data)
    elif type == 3:
        create_document_header(document, type, data, filepathString)
        #  create_application_framework(document)
        #  populate_application_framework(document, data)
        #  create_macro_framework(document)
        #  populate_macro_framework(document, data)
        populate_domain_framework(document, data)
    else:
        print("Error processing type")
    return document


def create_document_header(doc, type, data, filepathSting):
    """
    Creates the header for the document. This contains the main score and scores of the assessed criteria.
    """

    titleAddition = ""
    paraAdditionOne = ""
    paraAdditionTwo = ""
    if type == 1:
        titleAddition = "Application Hardening Settings"
        paraAdditionOne = "Application Hardening"
        paraAdditionTwo = "score"
    elif type == 2:
        titleAddition = "Microsoft Office Macro Settings"
        paraAdditionOne = "Microsoft Office Settings"
        paraAdditionTwo = "score"
    elif type == 3:
        titleAddition = "Application Hardening and Microsoft Office Macro Settings"
        paraAdditionOne = "Application Hardening and Microsoft Office Settings"
        paraAdditionTwo = "scores"
    else:
        print("Error matching Type")
        return
    doc.add_heading('Essential Eight ' + titleAddition + ' Report', 0)
    underHeaderPara = doc.add_paragraph()
    underHeaderPara.add_run(
        'This document contains the Essential Eight scores for Application Hardening and Microsoft' +
        ' Office Settings which were assessed from the provided file(s): \n' + filepathSting +
        '\n\nIt will provide an overview of the overall scores before providing reasoning based upon' +
        ' the GPOs in the file(s) which were provided.')


def save_document(doc, path):
    """Saves the document to the path specified. Path is RELATIVE to current directory is full directory is not used."""

    doc.save(path)


def populate_Overall_Score_framework(doc, data):
    applicationHardeningScore = 5;
    macroSettingsScore = 5;

    # score_list = [r.score for n, r in application_hardening_data if isinstance(r.score, int)]
    # if score_list:
    #    application_hardening_score = min([r.score for n, r in application_hardening_data if isinstance(r.score, int)])
    # else:
    #    application_hardening_score = "Not application"

    for gpo in data:
        if gpo.applicationHardening.return_score() < applicationHardeningScore:
            applicationHardeningScore = gpo.applicationHardening.return_score()
        if gpo.macroSettings.return_score() < macroSettingsScore:
            macroSettingsScore = gpo.macroSettings.return_score()

    if macroSettingsScore == 5:
        macroSettingsScore = 0
    if applicationHardeningScore == 5:
        applicationHardeningScore = 0

    doc.add_heading("Overall Score: " + str(min(applicationHardeningScore, macroSettingsScore)), level=2)
    paraOne = doc.add_paragraph()
    paraOne.add_run("   Application Hardening: " + str(applicationHardeningScore) + "\n")
    paraOne.add_run("   Microsoft Office Macro Settings: " + str(macroSettingsScore) + "\n")


def populate_domain_framework(doc, data):
    """
    creates a listing for all domains in the supplied data file
    """

    previous = ""
    for gpo in data:
        if previous != gpo.domain:
            doc.add_heading('Domain: ' + gpo.domain)
            populate_Overall_Score_framework(doc, data)
        doc.add_heading(gpo.GPOname + " overall score : " +
                        str(min(gpo.applicationHardening.return_score(), gpo.macroSettings.return_score())), level=2)
        create_application_framework(doc)
        populate_application_framework(doc, gpo)
        create_macro_framework(doc)
        populate_macro_framework(doc, gpo)
        previous = gpo.domain


def create_application_framework(doc):
    """Creates the basic framework before Application Hardening is assessed."""

    applicationScore1 = doc.add_heading('Application Hardening:', level=3)
    # paraOne = doc.add_paragraph()
    # paraOne.add_run('Overall Score: ' + '[insert score here] \n', style = 'FirstFont')
    # paraOne.add_run('Reasoning: ', style = 'FirstFont')


def populate_application_framework(doc, gpo):
    """
    Populates the document with each individual GPO and how they have been assessed against
    the Application Hardening framework.
    """

    #  for gpo in data:
    score = gpo.applicationHardening.return_score()
    para1 = doc.add_paragraph()
    para1.add_run('Score : ' + str(score) + '\n', style='FirstFont')
    para1.add_run('Reasoning: ', style='FirstFont')

    para2 = doc.add_paragraph()
    paragraph_format = para2.paragraph_format
    paragraph_format.left_indent = Inches(0.5)
    para2.add_run(gpo.application_to_string(), style='gpoFont')
    # print(gpo.GPOname)


def create_macro_framework(doc):
    """
    Creates the basic framework before macro's are assessed.
    """
    macroScore1 = doc.add_heading('Microsoft Office Macro Settings: ', level=3)
    #paraOne = doc.add_paragraph()
    #paraOne.add_run('Overall Score: ' + '[insert score here] \n', style='FirstFont')
    #paraOne.add_run('Reasoning: ', style='FirstFont')


def populate_macro_framework(doc, gpo):
    """
    Populates the document with each individual GPO and how they have been assessed against the Macro Office Settings
    framework.
    """
    # for gpo in data:
    score = gpo.macroSettings.return_score()
    para1 = doc.add_paragraph()
    para1.add_run('Score : ' + str(score) + '\n', style='FirstFont')
    para1.add_run('Reasoning: ', style='FirstFont')

    para2 = doc.add_paragraph()
    paragraph_format = para2.paragraph_format
    paragraph_format.left_indent = Inches(0.5)
    para2.add_run(gpo.macro_settings_to_string(), style='gpoFont')

    # print(gpo.GPOname)


def create_character_font(styles, size, font, name):
    """
    Creates a new character font and saves it to the stylesheet provided
    """

    newFont = styles.add_style(name, WD_STYLE_TYPE.CHARACTER)
    newFonts = newFont.font
    newFonts.size = Pt(size)
    newFonts.name = font


def adjust_font(doc, size, style, bold=False, italic=False, color=RGBColor(0, 0, 0)):
    """
    Adjusts Fonts which already exist within the documents stylesheet
    Defaulted Color is Black
    """
    style = doc.styles[style]
    styling = style.font
    styling.size = Pt(size)
    styling.bold = bold
    styling.italic = italic
    styling.color.rgb = color


def create_fonts(doc, styles):
    """Static creation of all the fonts used in the word document"""

    create_character_font(styles, 16, 'Calibri (Body)', 'FirstFont')
    create_character_font(styles, 14, 'Calibri (Body)', 'UnderHeadingsFont')
    create_character_font(styles, 11, 'Segoe UI', 'gpoFont')
    adjust_font(doc, 16, 'Heading 1', True)

# createEssentialEightDocument("test", 3)
